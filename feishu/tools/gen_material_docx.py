# -*- coding: utf-8 -*-
"""生成《飞书深度打通 · 技术选型与实施材料》Word 文档。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

ACCENT = RGBColor(0x2E, 0x5B, 0xFF)   # 飞书蓝
GREY = RGBColor(0x66, 0x66, 0x66)
DARK = RGBColor(0x22, 0x22, 0x22)

doc = Document()


def _cjk(run):
    run.font.name = "Microsoft YaHei"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")

# 中文字体
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E5BFF")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rPr.append(rFonts)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def body(text, size=10.5, color=DARK, space_after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def h1(text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.color.rgb = ACCENT
    run.font.size = Pt(16)
    _cjk(run)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    run.font.color.rgb = DARK
    run.font.size = Pt(13)
    _cjk(run)
    return p


def repo_table(rows, headers, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _cjk(run)
        set_cell_bg(hdr[i], "2E5BFF")
    for row in rows:
        cells = table.add_row().cells
        # 列:仓库(带链接), 星标, 语言, 描述/用途, 推荐
        name, url, star, lang, desc, rec = row
        # 仓库名 + 链接
        p0 = cells[0].paragraphs[0]
        add_hyperlink(p0, url, name)
        for run in p0.runs:
            run.font.size = Pt(9)
        for idx, val in enumerate([star, lang, desc, rec], start=1):
            c = cells[idx]
            c.text = ""
            run = c.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            _cjk(run)
            if idx == 4:  # 推荐列
                if val.startswith("★★★"):
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                    run.bold = True
                elif val.startswith("★★"):
                    run.font.color.rgb = RGBColor(0xD3, 0x5B, 0x00)
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(w)
    return table


# ────────────────────────── 封面 ──────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(60)
r = title.add_run("售前方案产线 · 飞书深度打通")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = ACCENT
_cjk(r)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("开源选型调研 · 技术方案 · 实施材料")
r.font.size = Pt(14)
r.font.color.rgb = GREY
_cjk(r)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(24)
r = meta.add_run(
    f"调研日期:2026-06-23   |   数据来源:GitHub Search(星标为调研当日数值)\n"
    f"适用仓库:sales-proposal(售前方案产线)"
)
r.font.size = Pt(10)
r.font.color.rgb = GREY
_cjk(r)

doc.add_page_break()

# ───────────────────── 1. 背景与目标 ─────────────────────
h1("一、背景与目标")
body("本仓库是一套售前方案产线:一支角色化的 Agent 团队(7 个售前专家)+ 一套培养新人驾驭它的"
     "训练营教材,方案与教材均以 Markdown 为载体。本次目标是把这套产线与飞书深度打通,"
     "让产出能在飞书侧协作、播报与看板化管理。", space_after=8)
body("结合产线特征,最具价值的三条打通线为:", bold=True, space_after=4)
bullet("方案 / 教材 ↔ 飞书云文档、知识库的双向同步;")
bullet("销售管道 ↔ 多维表格(Bitable)看板;")
bullet("让 7 个 Agent 直接操作飞书(云文档 / 多维表格 / 消息)。")

# ───────────────────── 2. 三层架构 ─────────────────────
h1("二、总体方案:三层架构")
body("三层共用一份应用凭证,互不依赖、可单独启用。", space_after=8)

arch = [
    ("① Agent 层", "lark-openapi-mcp(官方 MCP)",
     "让 7 个 Agent 原生操作飞书云文档 / 多维表格 / 消息,在对话中即时落地。"),
    ("② 文档同步层", "feishu-cli(社区)",
     "方案与教材 Markdown ↔ 飞书云文档双向无损同步(40+ 块类型,含表格、代码、Mermaid)。"),
    ("③ 自动化层", "oapi-sdk-python(官方 SDK)",
     "方案状态 → 机器人卡片播报;销售管道 → 多维表格看板。脚本化、可进 CI、可被 Agent 调用。"),
]
t = doc.add_table(rows=1, cols=3)
t.style = "Light Grid Accent 1"
hdrs = ["层", "核心组件", "解决什么"]
for i, htext in enumerate(hdrs):
    c = t.rows[0].cells[i]
    c.text = ""
    run = c.paragraphs[0].add_run(htext)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _cjk(run)
    set_cell_bg(c, "2E5BFF")
for layer, comp, desc in arch:
    cells = t.add_row().cells
    for idx, val in enumerate([layer, comp, desc]):
        cells[idx].text = ""
        run = cells[idx].paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if idx == 0:
            run.bold = True
            run.font.color.rgb = ACCENT
        _cjk(run)
for i, w in enumerate([0.9, 2.0, 3.6]):
    for row in t.rows:
        row.cells[i].width = Inches(w)

# ───────────────────── 3. 开源仓库调研清单 ─────────────────────
doc.add_page_break()
h1("三、开源仓库调研清单")
body("以下为 GitHub 上飞书 / Lark 相关高评分仓库的完整调研结果,按"
     "「官方」「社区」「避坑」三组列出,含链接、星标、语言、用途与推荐度。", space_after=10)

h2("3.1 官方仓库(larksuite 组织,优先选)")
official = [
    ("larksuite/cli", "https://github.com/larksuite/cli", "14.6k", "Go",
     "官方 CLI,200+ 命令 + 20+ AI Agent Skills,覆盖 Docs/Base/Sheets/Messenger/Calendar/Mail/Tasks/Meetings 等。",
     "★★★ CLI/脚本主干"),
    ("larksuite/lark-openapi-mcp", "https://github.com/larksuite/lark-openapi-mcp", "742", "TS",
     "官方 OpenAPI MCP 服务(npm 包 @larksuiteoapi/lark-mcp),直接挂到 Claude Code / Cursor。",
     "★★★ 本方案①层"),
    ("larksuite/oapi-sdk-python", "https://github.com/larksuite/oapi-sdk-python", "525", "Python",
     "官方 Python SDK(包名 lark-oapi,v2_main 活跃),全量 OpenAPI + 事件回调。",
     "★★★ 本方案③层"),
    ("larksuite/oapi-sdk-go", "https://github.com/larksuite/oapi-sdk-go", "598", "Go",
     "官方 Go SDK(v3_main),全量 OpenAPI。", "★★ 备选(Go 栈)"),
    ("larksuite/node-sdk", "https://github.com/larksuite/node-sdk", "273", "TS",
     "官方 Node SDK,活跃维护,取代已废弃的 oapi-sdk-nodejs。", "★★ 备选(Node 栈)"),
    ("larksuite/oapi-sdk-java", "https://github.com/larksuite/oapi-sdk-java", "305", "Java",
     "官方 Java SDK(v2_main)。", "★ 备选(Java 栈)"),
    ("larksuite/openclaw-lark", "https://github.com/larksuite/openclaw-lark", "2.3k", "TS",
     "飞书官方出品的 OpenClaw 飞书 / Lark Channel 插件。", "★ 机器人接入参考"),
    ("larksuite/lark-samples", "https://github.com/larksuite/lark-samples", "85", "Python",
     "官方示例(发消息 / 卡片 / 多维表格等),入门参考。", "★ 示例参考"),
]
repo_table(official, ["仓库", "星标", "语言", "用途 / 描述", "推荐度"],
           [1.7, 0.5, 0.5, 3.0, 1.0])

h2("3.2 社区高分仓库(贴合「方案=Markdown」场景)")
community = [
    ("riba2534/feishu-cli", "https://github.com/riba2534/feishu-cli", "1.2k", "Go",
     "核心能力是 Markdown ↔ 飞书文档双向无损转换;含 doc/wiki/sheet/bitable/msg/calendar 等子命令。",
     "★★★ 本方案②层"),
    ("Wsine/feishu2md", "https://github.com/Wsine/feishu2md", "2.2k", "Go",
     "一键把飞书文档导出为 Markdown(注:作者在征集维护者)。", "★★ 导出兜底"),
    ("cso1z/Feishu-MCP", "https://github.com/cso1z/Feishu-MCP", "695", "TS",
     "飞书文档 + 任务的 MCP / CLI + Skill,主打无缝接 Claude Code / Cursor / Cline。",
     "★★ ①层社区替代"),
    ("chyroc/lark", "https://github.com/chyroc/lark", "472", "Go",
     "社区 Go SDK,覆盖全部 Open API + 事件回调,封装比官方更顺手。", "★★ Go 栈替代"),
    ("eternalfree/feishu-doc-export", "https://github.com/eternalfree/feishu-doc-export", "809", "C#",
     "飞书文档批量导出服务。", "★ 批量导出"),
    ("ConnectAI-E/feishu-openai", "https://github.com/ConnectAI-E/feishu-openai", "5.6k", "Go",
     "飞书 ×(GPT-4 + DALL·E + Whisper)机器人:语音对话、文档导出、表格分析等。", "★ AI 机器人参考"),
    ("zarazhangrui/lark-coding-agent-bridge", "https://github.com/zarazhangrui/lark-coding-agent-bridge",
     "1.4k", "TS",
     "把飞书 / Lark 与本地 Claude Code / Codex 桥接:流式卡片、按会话隔离、多工作区。", "★ Agent 桥接参考"),
    ("bestony/ChatGPT-Feishu", "https://github.com/bestony/ChatGPT-Feishu", "1.2k", "JS",
     "给飞书准备的 ChatGPT 机器人,轻量易部署。", "★ 机器人参考"),
]
repo_table(community, ["仓库", "星标", "语言", "用途 / 描述", "推荐度"],
           [1.7, 0.5, 0.5, 3.0, 1.0])

h2("3.3 避坑清单(已归档 / 废弃,勿用)")
deprecated = [
    ("larksuite/oapi-sdk-nodejs", "https://github.com/larksuite/oapi-sdk-nodejs", "81", "TS",
     "已 DEPRECATED,改用 larksuite/node-sdk。", "✗ 废弃"),
    ("larksuite/feishu", "https://github.com/larksuite/feishu", "86", "Python",
     "WIP 且已 archived 的早期 Python SDK,改用 oapi-sdk-python。", "✗ 归档"),
    ("larksuite/botframework-go", "https://github.com/larksuite/botframework-go", "107", "Go",
     "已废弃,改用 oapi-sdk-go。", "✗ 废弃"),
    ("larksuite/appframework-java", "https://github.com/larksuite/appframework-java", "39", "Java",
     "已废弃,改用 oapi-sdk-java。", "✗ 废弃"),
]
repo_table(deprecated, ["仓库", "星标", "语言", "说明", "状态"],
           [1.7, 0.5, 0.5, 3.0, 1.0])

# ───────────────────── 4. 选型结论 ─────────────────────
doc.add_page_break()
h1("四、选型结论与理由")
items = [
    ("① Agent 层 → larksuite/lark-openapi-mcp(官方)",
     "官方维护、跟随 OpenAPI 演进,稳定性最佳;以 MCP 形式直接挂进 Claude Code,"
     "让 7 个 Agent 在对话中原生读写飞书,是「深度打通」价值最高的一环。"),
    ("② 文档同步层 → riba2534/feishu-cli(社区)",
     "社区 1.2k 星,核心能力正是 Markdown ↔ 云文档 40+ 块类型无损双向转换,"
     "与本仓库「方案 / 教材即 Markdown、交叉引用走相对路径」的形态完全契合;可脚本化、可进 CI。"),
    ("③ 自动化层 → larksuite/oapi-sdk-python(官方)",
     "官方 Python SDK,与本仓库文档调性、新人上手成本最匹配;用于实现确定性的"
     "状态卡片播报与多维表格看板,可被 Agent 直接当工具调用。"),
]
for head, desc in items:
    body(head, bold=True, color=ACCENT, space_after=2)
    body(desc, space_after=10)

body("分工原则:", bold=True, space_after=2)
bullet("文档类产出走第②层(无损、留痕、进 Git);", )
bullet("结构化 / 状态类产出走第③层(看板、播报);")
bullet("对话中即时、零散的操作走第①层 MCP。")

# ───────────────────── 5. 已交付实现 ─────────────────────
h1("五、已交付实现(本仓库)")
body("已在分支 claude/optimistic-cerf-0ydljr 完成三层落地,目录结构如下:",
     space_after=8)
deliver = [
    ("根 .mcp.json", "内置官方 lark-mcp,env 占位凭证,预设 doc/bitable/im 工具集。"),
    ("feishu/README.md", "三层总览、快速开始 5 步、安全约定。"),
    ("feishu/.env.example", "三层共用凭证模板(.env 已 gitignore)。"),
    ("feishu/01-agent-mcp/README.md", "MCP 接入、应用态/用户态、最小权限、Agent 提示词示例。"),
    ("feishu/02-doc-sync/", "feishu-cli 用法 + sync.sh 批量同步脚本 + sync-map 映射模板。"),
    ("feishu/03-automation/", "lark_client.py(鉴权)、notify.py(卡片播报)、pipeline_base.py(多维表格看板)、requirements.txt、fields.example.json。"),
    ("curriculum/05-飞书打通-使用手册.md", "新人视角:Agent → 飞书能力对照表、嵌入五段流程、四个即用动作。"),
    (".github/workflows/feishu-sync.yml", "定时把教材同步到飞书知识库的 CI(详见第六章)。"),
    ("README / CLAUDE.md / 00-README", "同步更新交叉引用与维护约定。"),
]
for path, desc in deliver:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(path + " — ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(desc)

# ───────────────────── 6. 定时同步 CI ─────────────────────
doc.add_page_break()
h1("六、定时同步 CI(GitHub Actions)")
body("已新增工作流 .github/workflows/feishu-sync.yml,把「教材目录定时 sync.sh push 到飞书"
     "知识库」做成 CI,让知识库内容自动保鲜。", space_after=8)

body("触发与行为", bold=True, color=ACCENT, space_after=2)
bullet("触发:每周一 09:00(北京时间)定时 + 支持手动触发(可指定只同步某个分组)。")
bullet("幂等:以 SYNC_UPDATE_ONLY=1 运行,只更新映射表里【已有 doc_id】的条目,绝不新建文档,"
       "避免每次定时跑都往知识库塞重复文档。", )
bullet("流程:Checkout → 装 Go/feishu-cli → 装 Python/PyYAML → 由 Secret 生成 sync-map.yaml "
       "→ 运行 sync.sh push。")
bullet("守卫:未配置凭证时自动跳过(notice 提示),不会让 CI 报红。")

body("需要在仓库配置的 Secrets / Variables", bold=True, color=ACCENT, space_after=2)
ci_cfg = [
    ("Secret  FEISHU_APP_ID", "自建应用 App ID。"),
    ("Secret  FEISHU_APP_SECRET", "自建应用 App Secret。"),
    ("Secret  FEISHU_SYNC_MAP", "sync-map.yaml 的完整内容(含各教材文件 ↔ doc_id)。因含具体 id 不入库,故走 Secret。"),
    ("Variable  FEISHU_DOMAIN", "可选;cn(默认)或 intl。"),
]
for k, v in ci_cfg:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(k + " — ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(v)

body("首次入库(一次性人工操作)", bold=True, color=ACCENT, space_after=2)
body("feishu-cli 的知识库无直接 import 命令;把一篇 Markdown 首次放进知识库的标准做法是"
     "「先建文档、再移入知识库空间」,之后 CI 便能按 doc_id 自动更新其内容:", space_after=4)
code = doc.add_paragraph()
code.paragraph_format.left_indent = Inches(0.3)
cr = code.add_run(
    "feishu-cli doc import <file.md> --title <标题>          # 建文档,记下 doc token\n"
    "feishu-cli wiki spaces                                  # 取目标知识库 space_id\n"
    "feishu-cli wiki move-docs <doc_token> --space-id <id>   # 移入知识库\n"
    "# 然后把 doc_id 写进 sync-map(即 FEISHU_SYNC_MAP),CI 自动保鲜"
)
cr.font.name = "Consolas"
cr.font.size = Pt(9)
cr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ───────────────────── 7. 接手须知 ─────────────────────
doc.add_page_break()
h1("七、你接手后要做的(凭证类,我无法代办)")
body("代码与文档已全部就位;下面这些涉及飞书账号 / 凭证 / 群与表的人工动作,需要你来完成。",
     space_after=8)

handover = [
    ("① 建应用、配凭证",
     ["在飞书开放平台创建企业自建应用,拿到 App ID / App Secret;",
      "执行 cp feishu/.env.example feishu/.env,把 App ID / App Secret 填进去;",
      "（CI 还需把这两项配成仓库 Secret:FEISHU_APP_ID / FEISHU_APP_SECRET。）"]),
    ("② 开权限、发版本",
     ["按各层 README 列出的最小权限集,在应用后台开通对应 scope;",
      "改完权限必须在开放平台发布新版本才生效——这一步最容易漏。"]),
    ("③ 拉群、建表、回填 id",
     ["把应用机器人拉进目标通知群,用 feishu-cli chat list 取 chat_id;",
      "在飞书手建销售管道多维表格(字段照 fields.example.json),取 app_token / table_id;",
      "把 chat_id / app_token / table_id 回填进 feishu/.env(对应 FEISHU_NOTIFY_CHAT_ID / "
      "FEISHU_PIPELINE_APP_TOKEN / FEISHU_PIPELINE_TABLE_ID)。"]),
    ("④ 启用 CI 定时同步(可选)",
     ["把教材首次入库(见第六章一次性操作),把含 doc_id 的 sync-map 内容配成 Secret FEISHU_SYNC_MAP;",
      "之后每周一 CI 自动把教材推到知识库,无需再管。"]),
]
for head, subs in handover:
    body(head, bold=True, color=ACCENT, space_after=3)
    for s in subs:
        bullet(s)
    body("", space_after=2)

body("安全与避坑", bold=True, color=ACCENT, space_after=2)
bullet("凭证只进 .env / GitHub Secret,永不入库(已 gitignore);仓库内 *.example.* 均为脱敏模板。")
bullet("改了权限 scope 必须在开放平台发布新版本才生效。")
bullet("机器人未进群直接发消息必然失败,先拉群再取 chat_id。")
bullet("CI 默认 update-only:sync-map 里没填 doc_id 的条目不会被同步,首次入库需人工一次。")

# 页脚
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("售前方案产线 · 飞书深度打通 · 技术选型与实施材料")
fr.font.size = Pt(8)
fr.font.color.rgb = GREY
_cjk(fr)

out = "/home/user/sales-proposal/飞书深度打通-技术选型与实施材料.docx"
doc.save(out)
print("saved:", out)
